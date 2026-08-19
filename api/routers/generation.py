from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import StreamingResponse
from typing import Optional, List
import pandas as pd
import numpy as np
import io
from docx import Document
from datetime import date

from api.deps import get_db, DataLoader
from api.schemas import MessageRequest, MessageResponse, FactUsed, ReportRequest, ReportResponse
from api.llm import generate_with_numeric_validation

router = APIRouter(prefix="/api", tags=["Generation"])

@router.post("/message", response_model=MessageResponse)
def post_message(req: MessageRequest):
    db = get_db()
    df = db.expectation
    
    # 1. 세그먼트 데이터 조회 (최신 연도 기준)
    latest_year = int(df['연도'].max())
    seg_row = df[
        (df['키'] == req.region) &
        (df['성별'] == req.sex) &
        (df['연령'] == req.age) &
        (df['암종'] == req.cancer) &
        (df['연도'] == latest_year)
    ]
    if seg_row.empty:
        raise HTTPException(status_code=404, detail="해당 조건의 세그먼트 데이터를 찾을 수 없습니다.")
        
    row = seg_row.iloc[0]
    our_rate = float(row['수검률'])
    
    # 2. 전국 평균 데이터 조회
    df_nat = db.national_avg
    latest_nat_year = int(df_nat['연도'].max()) if not df_nat.empty else latest_year
    nat_row = df_nat[
        (df_nat['성별'] == req.sex) &
        (df_nat['연령'] == req.age) &
        (df_nat['암종'] == req.cancer) &
        (df_nat['연도'] == latest_nat_year)
    ]
    nat_rate = float(nat_row.iloc[0]['전국_평균']) if not nat_row.empty else our_rate
    
    # 3. 암종별 의학적 팩트 구성
    cancer_periods = {
        "위암": "2년 (만 40세 이상)",
        "대장암": "1년 (만 50세 이상)",
        "간암": "6개월 (만 40세 이상 고위험군)",
        "유방암": "2년 (만 40세 이상 여성)",
        "자궁경부암": "2년 (만 20세 이상 여성)",
        "폐암": "2년 (만 54~74세 중 고위험군)"
    }
    
    cancer_supplies = {
        "대장암": "채변통 (보건소·보건지소·인근 약국에서 사전 수령 가능)",
        "위암": "내시경 예약 (지정 검진기관)",
        "간암": "초음파 및 혈액검사 예약 (지정 검진기관)",
        "유방암": "유방촬영 예약 (지정 검진기관)",
        "자궁경부암": "자궁경부세포검사 예약 (지정 검진기관)",
        "폐암": "저선량 CT 예약 (지정 검진기관)"
    }
    
    period = cancer_periods.get(req.cancer, "2년")
    supply_info = cancer_supplies.get(req.cancer, "지정 검진기관 예약")
    
    region_display = req.region.split('|')[1] if '|' in req.region else req.region
    
    facts = [
        FactUsed(label="검진 대상", value=f"{req.sex} {req.age}세"),
        FactUsed(label="검진 주기", value=period),
        FactUsed(label="본인부담", value="건강보험료 하위 50% 무료"),
        FactUsed(label="준비 사항 및 방법", value=supply_info)
    ]
    
    # 4. LLM 프롬프트 구성 (수치 정보 제공)
    our_rate_str = f"{our_rate:.2f}%"
    nat_rate_str = f"{nat_rate:.2f}%"
    
    prompt = f"""
    아래 정보를 바탕으로, {region_display} 보건소에서 {req.sex} {req.age}세 주민들을 대상으로 발송할 '{req.cancer}' 검진 독려 안내 문자 메시지 초안을 작성하십시오.
    
    [기본 정보]
    - 발신처: {region_display}보건소
    - 대상: {req.sex} {req.age}세
    - 대상 암종: {req.cancer}
    - 검진 주기: {period}
    - 준비 사항 및 방법: {supply_info}
    - 본인 부담금: 건강보험료 하위 50% 무료 (또는 의료급여수급자 무료)
    
    [통계 정보]
    - {region_display}의 해당 집단 실제 검진 수검률: {our_rate_str}
    - 전국 평균 수검률: {nat_rate_str}
    
    [요청 톤]
    - {req.tone} (standard: 표준적이고 공손함, polite: 정중하고 따뜻함, urgent: 수검 필요성을 다소 강조함)
    
    [안내문 작성 원칙]
    - 반드시 {region_display}보건소 명의로 작성해야 합니다. 절대 '온코케어'라는 단어를 본문에 노출하지 마십시오.
    - 「암일 수 있다」류의 공포감을 유발하는 표현은 금지합니다.
    - 수검률을 나타내는 숫자는 오직 지문에 제공된 정확한 숫자({our_rate_str}, {nat_rate_str})만 사용하여 인용하십시오. 임의의 수치나 통계를 지어내서 기재하면 안 됩니다.
    - 공손한 어조로 작성하되, 문자 메시지 형태로 간결하고 가독성 좋게 작성하십시오.
    """
    
    system_instruction = (
        "당신은 보건소 담당자로서 주민들에게 친절하고 명확한 국가암검진 독려 안내 문자를 작성하는 AI입니다. "
        "반드시 지문에 제공된 정확한 수치 정보만 기재해야 하며, 제공되지 않은 임의의 숫자를 지어내어 생성해서는 절대 안 됩니다. "
        "또한 본문에 '온코케어' 명칭을 사용하지 마십시오."
    )
    
    generated_text = generate_with_numeric_validation(
        prompt=prompt,
        system_instruction=system_instruction
    )
    
    reflected = [
        f"{req.cancer} 검진 주기({period}) 명시 및 안내",
        "비용 안내 (하위 50% 무료 대상 인지 유도)",
        f"{req.tone} 톤에 따른 맞춤형 설득 문구 적용"
    ]
    if req.cancer == "대장암":
        reflected.append("채변통 사전 수령 장소 안내 (약국, 보건소 등)")
        
    return MessageResponse(
        text=generated_text,
        facts_used=facts,
        reflected=reflected
    )

@router.post("/report")
def post_report(req: ReportRequest):
    db = get_db()
    df = db.expectation
    
    # 1. 해당 지역 및 연도 데이터 필터링
    df_year = df[(df['연도'] == req.year) & (df['키'] == req.region)].copy()
    if df_year.empty:
        raise HTTPException(status_code=404, detail="해당 조건의 데이터를 찾을 수 없습니다.")
        
    # 요약 통계 계산
    target_sum = int(df_year['대상자'].sum())
    screened_sum = int(df_year['수검자'].sum())
    total_rate = screened_sum / target_sum * 100 if target_sum > 0 else 0.0
    
    # 가장 부진한 암종과 회복여지 상위 1위 세그먼트 파악
    # 필터: 연령n < 11 (75세 미만), 회복여지 >= 300
    p_df = df_year[(df_year['연령n'] < 11) & (df_year['회복여지'] >= 300)].copy()
    if not p_df.empty:
        p_df = p_df.sort_values(by='회복여지', ascending=False)
        top_seg = p_df.iloc[0]
        worst_cancer = top_seg['암종']
        worst_seg_info = f"{top_seg['성별']} {top_seg['연령']} ({worst_cancer})"
        recovery_limit = int(np.round(top_seg['회복여지']))
        our_worst_rate = float(top_seg['수검률'])
        exp_worst_rate = float(top_seg['기대치'])
        gap_worst = float(top_seg['잔차'])
    else:
        worst_cancer = "대장암"
        worst_seg_info = "미특정"
        recovery_limit = 0
        our_worst_rate = 0.0
        exp_worst_rate = 0.0
        gap_worst = 0.0
        
    # 전국 암종별 평균 대비 격차 확인
    df_nat = db.national_avg
    df_nat_year = df_nat[df_nat['연도'] == req.year]
    nat_cancer_rates = df_nat_year.groupby('암종')['전국_평균'].mean().to_dict()
    
    region_display = req.region.split('|')[1] if '|' in req.region else req.region
    
    # 2. LLM을 활용하여 각 단락 본문 생성
    # 팩트 수치들
    total_rate_str = f"{total_rate:.2f}%"
    our_worst_rate_str = f"{our_worst_rate:.2f}%"
    exp_worst_rate_str = f"{exp_worst_rate:.2f}%"
    gap_worst_str = f"{gap_worst:.2f}%p"
    recovery_limit_str = f"{recovery_limit}명"
    
    prompt = f"""
    {region_display}의 {req.year}년도 국가암검진 결과 분석 및 보건사업 개선 계획 보고서 작성을 위해 아래 사실들을 기반으로 다음 4가지 섹션을 작성하십시오.
    
    [통계적 사실]
    - 대상 지역: {region_display}
    - 기준 연도: {req.year}년
    - 전체 수검자 수: {screened_sum}명 (총 대상자: {target_sum}명)
    - {region_display} 전체 6대암 통합 수검률: {total_rate_str}
    - 가장 시급한 부진 세그먼트: {worst_seg_info}
    - 해당 부진 세그먼트의 실제 수검률: {our_worst_rate_str} (여건 통제 기반 기대 수검률: {exp_worst_rate_str})
    - 기대치 대비 격차: {gap_worst_str}
    - 해당 세그먼트가 기대치에 도달할 경우 추가 검진 달성 가능한 '회복 여지': {recovery_limit_str}
    
    [섹션 구성 내용]
    1. 성과지표 (통합 수검률 및 수치 위주의 현황 요약)
    2. 수행내용 (주민 대상 검진 사업 추진 성과 요약)
    3. 미달사유 (기대치 대비 특정 세그먼트 부진의 데이터 분석 근거 제시)
    4. 개선계획 (회복 여지를 타깃으로 한 차년도 개입 전략 방향)
    
    [작성 규칙]
    - 오직 위에 제공된 수치({total_rate_str}, {our_worst_rate_str}, {exp_worst_rate_str}, {gap_worst_str}, {recovery_limit_str}, {screened_sum}명, {target_sum}명)만 기재하고, 임의의 수치를 조작하거나 날조하여 작성하지 마십시오.
    - 공적 문서 톤앤매너(하십시오체, 개조식 및 정돈된 단락)를 적용하십시오.
    - 본문에 '온코케어' 명칭을 언급하지 마십시오.
    - 각 섹션의 결과는 가독성 높은 텍스트로 줄바꿈을 포함해 작성하십시오.
    """
    
    system_instruction = (
        "당신은 지자체 보건소 국가암검진 사업 결과 보고서를 작성하는 전문 행정 연구원 AI입니다. "
        "반드시 지문에 제시된 통계 및 수치 정보만을 반영하여 사실에 기반한 보고서를 작성해야 합니다. "
        "임의의 다른 숫자나 연도를 추가하지 마십시오."
    )
    
    report_text = generate_with_numeric_validation(
        prompt=prompt,
        system_instruction=system_instruction
    )
    
    sections = [
        "성과지표 - 통합 검진 실적 요약",
        "수행내용 - 추진 현황",
        "미달사유 - 세그먼트별 격차 분석",
        "개선계획 - 중점 관리 및 전략 방안"
    ]
    
    # 3. 포맷별 응답 처리
    if req.format == "docx":
        # python-docx를 사용하여 실제 Word 파일 빌드
        doc = Document()
        doc.add_heading(f"{region_display} {req.year}년 국가암검진 결과 보고서", 0)
        
        doc.add_heading("1. 개요 및 성과지표", level=1)
        doc.add_paragraph(f"본 보고서는 {region_display}의 {req.year}년도 국가암검진 통합 실적을 여건 통제 기대치 모델 및 유사 지자체 비교 기법을 기반으로 분석하여 작성되었습니다.")
        
        # 텍스트에서 4가지 단락으로 나뉘어 있으므로 전체 텍스트를 줄바꿈 기준으로 추가
        # 통째로 추가하되 섹션으로 분리
        doc.add_heading("2. 상세 분석 및 개선 계획 본문", level=1)
        
        for p in report_text.split('\n\n'):
            if p.strip():
                doc.add_paragraph(p.strip())
                
        # 하단 필수 출처 및 경고 문구 추가
        doc.add_paragraph("\n\n* 데이터 출처 — 국민건강보험공단 / 국민연금공단")
        doc.add_paragraph("* 이 값은 평가가 아니라 검토 시작점입니다.")
        
        # 바이트 스트림 저장
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        filename = f"OncoCare_Report_{region_display}_{req.year}.docx"
        # 한글 파일명 헤더 대응
        from urllib.parse import quote
        headers = {
            'Content-Disposition': f"attachment; filename*=UTF-8''{quote(filename)}"
        }
        return StreamingResponse(file_stream, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers=headers)
        
    else:
        # JSON 응답 반환
        return ReportResponse(
            text=report_text,
            sections=sections
        )
